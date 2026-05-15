from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any, Dict
import base64
import io

def _format_paragraph(doc: Document, p_data: Dict[str, Any], rules: Dict[str, Any]):
    label = p_data.get("label", "BODY")
    text = p_data['data']["text"]

    para = doc.add_paragraph()
    
    # Style based on label
    if label == "TITLE":
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(rules.get("font_size_heading", 14))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif label == "AUTHORS":
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12) + 1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif label == "ABSTRACT":
        run = para.add_run("Abstract — ")
        run.bold = True
        para.add_run(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif label == "HEADING1":
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(rules.get("font_size_body", 12) + 2)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif label == "HEADING2":
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(rules.get("font_size_body", 12) + 1)
    elif label == "REFERENCES":
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12) - 2)
    else: # BODY
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    # Common formatting
    for run in para.runs:
        run.font.name = rules.get("font_family", "Times New Roman")
        
    para.paragraph_format.line_spacing = rules.get("line_spacing", 1.15)

def _format_table(doc: Document, t_data: Dict[str, Any], rules: Dict[str, Any]):
    table_data = t_data['data']
    
    if not table_data:
        return

    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # A default style

    for i, row_data in enumerate(table_data):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text

def _format_image(doc: Document, i_data: Dict[str, Any], rules: Dict[str, Any]):
    image_b64 = i_data['data']['content']
    
    try:
        image_data = base64.b64decode(image_b64)
        image_stream = io.BytesIO(image_data)
        doc.add_picture(image_stream, width=Inches(rules.get("image_width", 5.0)))
    except Exception as e:
        # If there is an error, add a placeholder text
        doc.add_paragraph(f"[Image placeholder: failed to load image - {e}]")

def format_manuscript(
    classified_data: dict,
    publication: str,
    doc_type: str,
    output_path: str,
    rules: dict
):
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    margins = rules.get("margins", {})
    section.top_margin = Inches(margins.get("top", 1))
    section.bottom_margin = Inches(margins.get("bottom", 1))
    section.left_margin = Inches(margins.get("left", 1))
    section.right_margin = Inches(margins.get("right", 1))
    
    if rules.get("columns") == 2:
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '2')
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '720')

    elements = classified_data.get("elements", [])
    
    for element in elements:
        element_type = element.get("type")
        label = element.get("label")

        if element_type == "paragraph":
            _format_paragraph(doc, element, rules)
        elif element_type == "table":
            _format_table(doc, element, rules)
            # Add a caption placeholder after the table
            doc.add_paragraph(f"[Table {label} placeholder caption]", style='Caption')
        elif element_type == "image":
            _format_image(doc, element, rules)
            # Add a caption placeholder after the image
            doc.add_paragraph(f"[Figure {label} placeholder caption]", style='Caption')
            
    doc.save(output_path)