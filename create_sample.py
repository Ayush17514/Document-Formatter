from docx import Document

def create_sample():
    doc = Document()
    doc.add_paragraph("A COMPREHENSIVE STUDY OF ARCHITECTURAL PATTERNS IN DECENTRALIZED SYSTEMS")
    doc.add_paragraph("Alice B. Researcher and Bob C. Engineer")
    doc.add_paragraph("University of Technology, Silicon Valley, CA")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This paper explores various architectural patterns used in modern decentralized systems. We analyze the trade-offs between consistency, availability, and partition tolerance (CAP theorem) and evaluate the performance of different consensus algorithms in large-scale deployments.")
    doc.add_paragraph("1. INTRODUCTION")
    doc.add_paragraph("Decentralized systems have gained significant traction in recent years due to the rise of blockchain technology and edge computing. These systems offer enhanced resilience and transparency compared to traditional centralized architectures.")
    doc.add_paragraph("1.1 Background")
    doc.add_paragraph("The concept of decentralization is not new, but modern implementations benefit from advanced networking and cryptography.")
    doc.add_paragraph("2. METHODOLOGY")
    doc.add_paragraph("We conducted a series of benchmarks across three different network topologies. Our metrics include latency, throughput, and fault tolerance threshold.")
    doc.add_paragraph("REFERENCES")
    doc.add_paragraph("[1] Lamport, L. (1978). Time, clocks, and the ordering of events in a distributed system.")
    doc.add_paragraph("[2] Nakamoto, S. (2008). Bitcoin: A Peer-to-Peer Electronic Cash System.")
    
    doc.save("sample_manuscript.docx")
    print("Created sample_manuscript.docx")

if __name__ == "__main__":
    create_sample()
