#!/usr/bin/env python
import os
import shutil
import uuid
import mammoth
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import re
from typing import Any, Dict, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io
import zipfile
from lxml import etree
import json
import google.generativeai as genai

# Logging configuration
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
LABEL_TITLE = "TITLE"
LABEL_AUTHORS = "AUTHORS"
LABEL_ABSTRACT = "ABSTRACT"
LABEL_HEADING1 = "HEADING1"
LABEL_HEADING2 = "HEADING2"
LABEL_BODY = "BODY"
LABEL_REFERENCES = "REFERENCES"
LABEL_TABLE = "TABLE"
LABEL_IMAGE = "IMAGE"

# --- Publication Rules ---

PUBLICATION_RULES = {
  "Research Paper": {
    "IEEE Access": {
      "font_family": "Times New Roman",
      "font_size_body": 10,
      "font_size_heading": 18,
      "columns": 2,
      "line_spacing": 1.0,
      "margins": { "top": 0.75, "bottom": 1.0, "left": 0.625, "right": 0.625 },
      "alignment": "JUSTIFIED"
    }
  }
}

DEFAULT_RULES = {
  "font_family": "Times New Roman",
  "font_size_body": 12,
  "font_size_heading": 14,
  "columns": 1,
  "line_spacing": 1.15,
  "margins": { "top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0 },
  "alignment": "JUSTIFIED"
}

# --- AI Integration ---

def get_gemini_client():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-1.5-flash')

async def resolve_formatting_rules(publication: str, doc_type: str) -> dict:
    if doc_type in PUBLICATION_RULES and publication in PUBLICATION_RULES[doc_type]:
        logger.info(f"Using strict rules for {publication}")
        final_rules = DEFAULT_RULES.copy()
        final_rules.update(PUBLICATION_RULES[doc_type][publication])
        return final_rules
    
    client = get_gemini_client()
    if not client:
        logger.warning("No Gemini API key. Using default rules.")
        return DEFAULT_RULES
        
    prompt = f'''
    You are an expert in academic publication formatting.
    Provide the official formatting rules for:
    Publication: {publication}
    Document Type: {doc_type}
    
    Return ONLY a JSON object with these keys:
    {{
        "font_family": "string",
        "font_size_body": "number",
        "font_size_heading": "number",
        "columns": "1 or 2",
        "line_spacing": "number",
        "margins": {{"top": "num", "bottom": "num", "left": "num", "right": "num"}}
    }}
    Use inches for margins. If unknown, use reasonable standard values.
    '''
    
    try:
        response = client.generate_content(prompt)
        text = response.text.strip()
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
            
        ai_rules = json.loads(text)
        
        final_rules = DEFAULT_RULES.copy()
        final_rules.update(ai_rules)
        
        if 'margins' in ai_rules and isinstance(ai_rules['margins'], dict):
            final_rules['margins'] = DEFAULT_RULES['margins'].copy()
            final_rules['margins'].update(ai_rules['margins'])

        logger.info(f"AI generated rules for {publication}")
        return final_rules
        
    except Exception as e:
        logger.error(f"AI rule resolution failed: {e}. Falling back to default rules.")
        return DEFAULT_RULES

def get_optimized_layout(elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    client = get_gemini_client()
    if not client:
        logger.warning("No Gemini API key found. Skipping layout optimization.")
        return elements

    text_content = "\n".join([e['data']['text'] for e in elements if e.get('type') == 'paragraph' and e.get('data')])
    
    prompt = f'''
    You are an expert in document layout and formatting.
    I have a document with the following text content:
    ---
    {text_content}
    ---
    
    I also have a list of elements (paragraphs, tables, images) in the order they appear in the original document.
    Your task is to re-order these elements to create a professional, readable document.
    
    Please provide the new order of elements as a comma-separated list of their original indices.
    For example: 0,1,3,2,4
    
    Original elements (with their index):
    {[(i, e.get('type')) for i, e in enumerate(elements)]}
    '''

    try:
        response = client.generate_content(prompt)
        new_order_str = response.text.strip()
        
        new_order = []
        for token in new_order_str.split(','):
            token = token.strip()
            if token == '':
                continue
            try:
                new_order.append(int(token))
            except ValueError:
                logger.warning(f"Skipping invalid index from AI: {token}")
                continue
        
        if len(new_order) != len(elements):
            logger.warning("AI returned an invalid new order. Using original order.")
            return elements
        
        reordered_elements = [elements[i] for i in new_order]
        
        logger.info("Successfully reordered elements with AI.")
        return reordered_elements

    except Exception as e:
        logger.error(f"AI layout optimization failed: {e}. Using original order.")
        return elements

# --- Parser ---

def _extract_text_from_docx_xml(filepath: str) -> List[str]:
    texts: List[str] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            for name in z.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    raw = z.read(name)
                    root = etree.fromstring(raw)
                    for t in root.findall(".//w:t", namespaces=ns):
                        if t is not None and t.text and t.text.strip():
                            texts.append(t.text.strip())
                except Exception:
                    continue
    except Exception:
        return []
    return texts

def parse_document(filepath: str) -> dict[str, Any]:
    try:
        doc = Document(filepath)
        elements: List[Dict[str, Any]] = []
        
        def add_element(element_type, data, index):
            elements.append({"type": element_type, "data": data, "index": index})

        for i, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            runs = []
            for run in para.runs:
                try:
                    size = run.font.size.pt if run.font and run.font.size else None
                except Exception:
                    size = None
                runs.append({
                    "text": run.text or "",
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "font_size": size,
                    "font_name": run.font.name or ""
                })
            
            paragraph_data = {
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": str(para.alignment) if getattr(para, 'alignment', None) else None,
                "runs": runs
            }
            add_element("paragraph", paragraph_data, i)

        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text or "")
                table_data.append(row_data)
            add_element("table", table_data, len(elements))

        for i, shape in enumerate(doc.inline_shapes):
            try:
                if "image" in shape.part.content_type:
                    image_data = base64.b64encode(shape.part.blob).decode('utf-8')
                    add_element("image", {"content": image_data, "content_type": shape.part.content_type}, len(elements))
            except Exception:
                continue

        if len(elements) == 0:
            xml_texts = _extract_text_from_docx_xml(filepath)
            if xml_texts:
                for i, t in enumerate(xml_texts):
                    elements.append({
                        "type": "paragraph",
                        "data": {"text": t, "style": "Unknown"},
                        "index": i
                    })
                logger.info(f"XML fallback recovered {len(xml_texts)} text nodes from docx: {filepath}")

        return {
            "elements": elements,
            "metadata": {
                "author": getattr(doc.core_properties, 'author', None),
                "title": getattr(doc.core_properties, 'title', None)
            }
        }

    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        raise

# --- Classifier ---

class DocumentClassifier:
    def __init__(self):
        self.in_references = False

    def classify(self, elements: list) -> list:
        classified_elements = []
        i = 0
        while i < len(elements):
            element = elements[i]
            if element.get("type") == "paragraph":
                j = i
                pseudo_table_buffer = []
                while j < len(elements) and elements[j].get("type") == "paragraph":
                    text = elements[j]["data"]["text"] if elements[j].get('data') else elements[j].get('text', '')
                    if len(text.split()) < 5 and text.strip() != "":
                        pseudo_table_buffer.append(elements[j])
                        j += 1
                    else:
                        break
                
                if len(pseudo_table_buffer) > 2:
                    table_data = [[p['data']['text'] if p.get('data') else p.get('text','')] for p in pseudo_table_buffer]
                    
                    new_table_element = {
                        "type": "table",
                        "data": table_data,
                        "index": pseudo_table_buffer[0]['index'] 
                    }
                    classified_elements.append(self._classify_table(new_table_element))
                    i = j
                    continue

            element_type = element.get("type")

            if element_type == "paragraph":
                classified_elements.append(self._classify_paragraph(element))
            elif element_type == "table":
                classified_elements.append(self._classify_table(element))
            elif element_type == "image":
                classified_elements.append(self._classify_image(element))
            else:
                classified_elements.append(element)
            i += 1
            
        return classified_elements

    def _classify_paragraph(self, element: Dict[str, Any]) -> Dict[str, Any]:
        text = element['data']['text'] if element.get('data') else element.get('text','')
        text_lower = text.lower()
        
        label = LABEL_BODY
        confidence = 0.7

        if self.in_references:
            label = LABEL_REFERENCES
            confidence = 0.9
        elif element.get("index") == 0 and len(text) < 200:
            label = LABEL_TITLE
            confidence = 0.8
        elif element.get("index", 999) < 5 and (re.search(r"@|\w+, \w+", text) or "university" in text_lower):
            label = LABEL_AUTHORS
            confidence = 0.8
        elif "abstract" in text_lower:
            label = LABEL_ABSTRACT
            confidence = 0.95
        elif text_lower.strip() in ["references", "bibliography", "works cited"]:
            label = LABEL_REFERENCES
            self.in_references = True
            confidence = 0.99
        elif len(text) < 100 and (re.match(r"^[I|V|X|\d]+\.", text) or text.isupper()):
            label = LABEL_HEADING1
            confidence = 0.8
        elif len(text) < 100 and re.match(r"^\d+\.\d+", text):
            label = LABEL_HEADING2
            confidence = 0.8

        element["label"] = label
        element["confidence"] = confidence
        return element

    def _classify_table(self, element: Dict[str, Any]) -> Dict[str, Any]:
        element["label"] = LABEL_TABLE
        element["confidence"] = 0.99
        return element

    def _classify_image(self, element: Dict[str, Any]) -> Dict[str, Any]:
        element["label"] = LABEL_IMAGE
        element["confidence"] = 0.99
        return element

def classify_document(parsed_data: dict) -> dict:
    classifier = DocumentClassifier()
    parsed_data["elements"] = classifier.classify(parsed_data.get("elements", []))
    return parsed_data

# --- Validator ---

def validate_document(classified_data: dict) -> dict:
    if isinstance(classified_data, list):
        paragraphs = classified_data
    else:
        paragraphs = (
            classified_data.get("paragraphs")
            or classified_data.get("classified")
            or classified_data.get("elements")
            or []
        )

    normalized = []
    for p in paragraphs:
        if isinstance(p, dict):
            normalized.append(p)
        elif isinstance(p, str):
            normalized.append({"data": {"text": p}, "label": "BODY"})
        else:
            continue

    paragraphs = normalized

    if not isinstance(paragraphs, list) or len(paragraphs) == 0:
        return {
            "score": 0,
            "issues": ["No paragraphs found in classified data"],
            "stats": {
                "total_paragraphs": 0,
                "label_distribution": {}
            }
        }

    labels = [p.get("label", "UNKNOWN") for p in paragraphs]

    issues = []
    score = 100

    if "TITLE" not in labels:
        issues.append("Missing Title")
        score -= 20
    if "ABSTRACT" not in labels:
        issues.append("Missing Abstract section")
        score -= 15
    if "REFERENCES" not in labels:
        issues.append("Missing References section")
        score -= 10

    counts = {}
    for l in labels:
        counts[l] = counts.get(l, 0) + 1

    return {
        "score": max(0, score),
        "issues": issues,
        "stats": {
            "total_paragraphs": len(paragraphs),
            "label_distribution": counts
        }
    }

# --- Formatter ---

def _format_paragraph(doc: Document, p_data: Dict[str, Any], rules: Dict[str, Any]):
    label = p_data.get("label", "BODY")
    text = p_data['data']["text"]

    para = doc.add_paragraph()
    
    if label == "TITLE":
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(rules.get("font_size_heading", 14))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif label == "AUTHORS":
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12) + 1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif label == "ABSTRACT":
        run = para.add_run("Abstract — ")
        run.bold = True
        para.add_run(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif label == "HEADING1":
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(rules.get("font_size_body", 12) + 2)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif label == "HEADING2":
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(rules.get("font_size_body", 12) + 1)
    elif label == "REFERENCES":
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12) - 2)
    else:
        run = para.add_run(text)
        run.font.size = Pt(rules.get("font_size_body", 12))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    for run in para.runs:
        run.font.name = rules.get("font_family", "Times New Roman")
        
    para.paragraph_format.line_spacing = rules.get("line_spacing", 1.15)

def _format_table(doc: Document, t_data: Dict[str, Any], rules: Dict[str, Any]):
    table_data = t_data['data']
    
    if not table_data:
        return

    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(table_data):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text

def _format_image(doc: Document, i_data: Dict[str, Any], rules: Dict[str, Any]):
    image_b64 = i_data['data']['content']
    
    try:
        image_data = base64.b64decode(image_b64)
        image_stream = io.BytesIO(image_data)
        doc.add_picture(image_stream, width=Inches(rules.get("image_width", 5.0)))
    except Exception as e:
        doc.add_paragraph(f"[Image placeholder: failed to load image - {e}]")

def format_manuscript(
    classified_data: dict,
    publication: str,
    doc_type: str,
    output_path: str,
    rules: dict
):
    doc = Document()
    
    section = doc.sections[0]
    margins = rules.get("margins", {})
    section.top_margin = Inches(margins.get("top", 1))
    section.bottom_margin = Inches(margins.get("bottom", 1))
    section.left_margin = Inches(margins.get("left", 1))
    section.right_margin = Inches(margins.get("right", 1))
    
    if rules.get("columns") == 2:
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '2')
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '720')

    elements = classified_data.get("elements", [])
    
    for element in elements:
        element_type = element.get("type")
        label = element.get("label")

        if element_type == "paragraph":
            _format_paragraph(doc, element, rules)
        elif element_type == "table":
            _format_table(doc, element, rules)
            doc.add_paragraph(f"[Table {label} placeholder caption]", style='Caption')
        elif element_type == "image":
            _format_image(doc, element, rules)
            doc.add_paragraph(f"[Figure {label} placeholder caption]", style='Caption')
            
    doc.save(output_path)

# --- FastAPI App ---

app = FastAPI(
    title="Manuscript Formatter AI",
    description="AI-powered manuscript formatting platform",
    version="3.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
OUTPUT_DIR = os.path.join(os.getcwd(), "outputs")
STATIC_DIR = os.path.join(BASE_DIR, "static")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

@app.get("/")
async def root():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))

@app.get("/api/options")
async def get_options():
    return JSONResponse(content=PUBLICATION_RULES)

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    file_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_DIR, f"{file_id}.docx")
    
    with open(filepath, "wb") as f:
        shutil.copyfileobj(file.file, f)
    
    try:
        parsed = parse_document(filepath)
        classified = classify_document(parsed)
        
        return {
            "file_id": file_id,
            "classified": classified["elements"]
        }
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail=f"Parsing failed: {str(e)}")

@app.post("/api/process")
async def process_file(payload: Dict[str, Any]):
    try:
        rules = await resolve_formatting_rules(payload["publication"], payload["doc_type"])
        
        optimized_elements = get_optimized_layout(payload["classified"])
        
        output_filename = f"formatted_{payload['file_id']}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        format_manuscript(
            classified_data={"elements": optimized_elements},
            publication=payload["publication"],
            doc_type=payload["doc_type"],
            output_path=output_path,
            rules=rules
        )
        
        with open(output_path, "rb") as doc_file:
            result = mammoth.convert_to_html(doc_file)
            preview_html = result.value
            
        return {
            "status": "success",
            "preview_html": preview_html,
            "rules": rules
        }
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download/{file_id}")
async def download_file(file_id: str):
    path = os.path.join(OUTPUT_DIR, f"formatted_{file_id}.docx")
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Processed file not found")
        
    return FileResponse(
        path, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"formatted_{file_id}.docx"
    )

@app.on_event("startup")
async def startup_event():
    """Clean up old files on startup."""
    for directory in [UPLOAD_DIR, OUTPUT_DIR]:
        for filename in os.listdir(directory):
            if os.path.isfile(os.path.join(directory, filename)):
                os.remove(os.path.join(directory, filename))
