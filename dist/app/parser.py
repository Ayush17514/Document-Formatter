from docx import Document
from typing import Any
from app.utils.logger import logger
import base64
from docx.document import Document as DocumentObject
from docx.table import Table

def parse_document(filepath: str) -> dict[str, Any]:
    try:
        doc = Document(filepath)
        
        elements = []
        
        # Helper to process and add elements
        def add_element(element_type, data, index):
            elements.append({"type": element_type, "data": data, "index": index})

        # Process paragraphs and their content
        for i, para in enumerate(doc.paragraphs):
            # Extract text and formatting
            text = para.text.strip()
            runs = []
            for run in para.runs:
                runs.append({
                    "text": run.text,
                    "bold": run.bold or False,
                    "italic": run.italic or False,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "font_name": run.font.name
                })
            
            paragraph_data = {
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": runs
            }
            add_element("paragraph", paragraph_data, i)

        # Process tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            add_element("table", table_data, len(elements))

        # Process images
        for i, shape in enumerate(doc.inline_shapes):
            if "image" in shape.part.content_type:
                image_data = base64.b64encode(shape.part.blob).decode('utf-8')
                add_element("image", {"content": image_data, "content_type": shape.part.content_type}, len(elements))

        return {
            "elements": elements,
            "metadata": {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title
            }
        }

    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        raise
