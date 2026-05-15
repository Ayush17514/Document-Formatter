from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from typing import Any

def format_manuscript(
    classified_data: dict,
    publication: str,
    doc_type: str,
    output_path: str,
    rules: dict
):
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    margins = rules.get("margins", {"top": 1, "bottom": 1, "left": 1, "right": 1})
    section.top_margin = Inches(margins["top"])
    section.bottom_margin = Inches(margins["bottom"])
    section.left_margin = Inches(margins["left"])
    section.right_margin = Inches(margins["right"])
    
    # Handle columns (Note: python-docx column support is basic, usually needs section manipulation)
    if rules.get("columns") == 2:
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '2')
        cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '720') # 0.5 inch
    
    paragraphs = classified_data["paragraphs"]
    
    for p_data in paragraphs:
        label = p_data["label"]
        text = p_data["text"]
        
        para = doc.add_paragraph()
        
        # Style based on label
        if label == "TITLE":
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(rules["font_size_heading"])
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif label == "AUTHORS":
            run = para.add_run(text)
            run.font.size = Pt(rules.get("font_size_body", 12) + 1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif label == "ABSTRACT":
            run = para.add_run("Abstract — ")
            run.bold = True
            para.add_run(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif label == "HEADING1":
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(rules.get("font_size_body", 12) + 2)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif label == "HEADING2":
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(rules.get("font_size_body", 12) + 1)
        elif label == "REFERENCES":
            run = para.add_run(text)
            run.font.size = Pt(rules.get("font_size_body", 12) - 2)
        else:
            # Body
            run = para.add_run(text)
            run.font.size = Pt(rules.get("font_size_body", 12))
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Common formatting
        for run in para.runs:
            run.font.name = rules.get("font_family", "Times New Roman")
            
        para.paragraph_format.line_spacing = rules.get("line_spacing", 1.15)
        
    doc.save(output_path)
