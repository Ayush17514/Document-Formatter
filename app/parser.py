from docx import Document
from typing import Any, List, Dict
from app.utils.logger import logger
import base64
from docx.document import Document as DocumentObject
from docx.table import Table
import zipfile
from lxml import etree


def _extract_text_from_docx_xml(filepath: str) -> List[str]:
    """Extract text nodes (w:t) from all .xml parts inside the .docx zip.
    This helps recover text placed in textboxes/shapes or other parts
    that python-docx.doc.paragraphs may miss.
    """
    texts: List[str] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            for name in z.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    raw = z.read(name)
                    root = etree.fromstring(raw)
                    for t in root.findall(".//w:t", namespaces=ns):
                        if t is not None and t.text and t.text.strip():
                            texts.append(t.text.strip())
                except Exception:
                    # ignore parsing issues for non-well-formed xml parts
                    continue
    except Exception:
        return []
    return texts


def parse_document(filepath: str) -> dict[str, Any]:
    try:
        doc = Document(filepath)
        
        elements: List[Dict[str, Any]] = []
        
        # Helper to process and add elements
        def add_element(element_type, data, index):
            elements.append({"type": element_type, "data": data, "index": index})

        # Process paragraphs and their content
        for i, para in enumerate(doc.paragraphs):
            # Extract text and formatting
            text = (para.text or "").strip()
            runs = []
            for run in para.runs:
                try:
                    size = run.font.size.pt if run.font and run.font.size else None
                except Exception:
                    size = None
                runs.append({
                    "text": run.text or "",
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "font_size": size,
                    "font_name": run.font.name or ""
                })
            
            paragraph_data = {
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": str(para.alignment) if getattr(para, 'alignment', None) else None,
                "runs": runs
            }
            add_element("paragraph", paragraph_data, i)

        # Process tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text or "")
                table_data.append(row_data)
            add_element("table", table_data, len(elements))

        # Process images
        for i, shape in enumerate(doc.inline_shapes):
            try:
                if "image" in shape.part.content_type:
                    image_data = base64.b64encode(shape.part.blob).decode('utf-8')
                    add_element("image", {"content": image_data, "content_type": shape.part.content_type}, len(elements))
            except Exception:
                continue

        # XML fallback: if no elements were recovered via python-docx,
        # attempt to pull text nodes from the .docx zip (covers textboxes/shapes)
        if len(elements) == 0:
            xml_texts = _extract_text_from_docx_xml(filepath)
            if xml_texts:
                for i, t in enumerate(xml_texts):
                    elements.append({
                        "type": "paragraph",
                        "data": {"text": t, "style": "Unknown"},
                        "index": i
                    })
                logger.info(f"XML fallback recovered {len(xml_texts)} text nodes from docx: {filepath}")

        return {
            "elements": elements,
            "metadata": {
                "author": getattr(doc.core_properties, 'author', None),
                "title": getattr(doc.core_properties, 'title', None)
            }
        }

    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        raise
